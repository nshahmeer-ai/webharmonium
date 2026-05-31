"""
WebHarmonium — Business & Implementation Roadmap DOCX Generator
Generates a professional, investor-ready 20-page document.
Run: python generate_docx.py
Output: WebHarmonium_Roadmap.docx (in d:/webharmonium/)
"""

import subprocess, sys, os

# ── Auto-install python-docx ──
def install(pkg):
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("Installing python-docx...")
    install("python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Color Constants (RGB) ──
C_PRIMARY    = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A
C_SECONDARY  = RGBColor(0x1E, 0x29, 0x3B)   # #1E293B
C_GOLD       = RGBColor(0xD4, 0xAF, 0x37)   # #D4AF37
C_CREAM      = RGBColor(0xF8, 0xF5, 0xEE)   # #F8F5EE
C_CTA        = RGBColor(0x22, 0xC5, 0x5E)   # #22C55E
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK_TEXT  = RGBColor(0x1E, 0x29, 0x3B)
C_GREY       = RGBColor(0x94, 0xA3, 0xB8)
C_LIGHT_GREY = RGBColor(0xF1, 0xF5, 0xF9)

# ── Helpers ──

def set_cell_bg(cell, hex_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hr(doc, color='D4AF37', thickness=2):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness * 4))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h1(doc, text, color=C_PRIMARY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def h2(doc, text, color=C_GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def h3(doc, text, color=C_SECONDARY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def body(doc, text, color=C_DARK_TEXT, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.line_spacing = Pt(16)
    return p

def bullet(doc, text, level=0, color=C_DARK_TEXT):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def label_pill(doc, text, color_hex='D4AF37'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'  {text.upper()}  ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_WHITE
    run.font.name = 'Calibri'
    # Can't do real pill in DOCX, use highlight
    run.font.highlight_color = None
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))
    return p

def page_break(doc):
    doc.add_page_break()

# ── Cover Page ──
def make_cover(doc):
    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run('🎵')
    r.font.size = Pt(48)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('WebHarmonium')
    r.bold = True
    r.font.size = Pt(38)
    r.font.color.rgb = C_PRIMARY
    r.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Business & Implementation Roadmap')
    r.font.size = Pt(18)
    r.font.color.rgb = C_GOLD
    r.font.name = 'Calibri'
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Zero Investment  →  SEO Growth  →  AdSense Revenue')
    r.font.size = Pt(14)
    r.font.color.rgb = C_GREY
    r.font.name = 'Calibri'

    doc.add_paragraph()
    add_hr(doc, 'D4AF37', 3)

    # Tagline box (table)
    t = doc.add_table(1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_bg(cell, '0F172A')
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(
        'South Asia\'s Best Free Online Harmonium Platform\n'
        'Build | Grow | Monetize — Starting at $0'
    )
    cr.bold = True
    cr.font.size = Pt(13)
    cr.font.color.rgb = C_CREAM
    cr.font.name = 'Calibri'
    cell.paragraphs[0].paragraph_format.space_before = Pt(12)
    cell.paragraphs[0].paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # Meta info
    info_pairs = [
        ('Document Type', 'Business & Technical Roadmap'),
        ('Version', '1.0 — Phase 1 MVP'),
        ('Platform', 'HTML · CSS · JavaScript · Netlify'),
        ('Initial Investment', '$0 (Zero Cost Launch)'),
        ('Target Market', 'South Asia + Global Diaspora'),
        ('Revenue Model', 'Google AdSense + Affiliate + Premium'),
        ('Prepared By', 'WebHarmonium Team'),
        ('Date', '2024'),
    ]
    t2 = doc.add_table(len(info_pairs), 2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.style = 'Table Grid'
    for i, (k, v) in enumerate(info_pairs):
        row = t2.rows[i]
        kc = row.cells[0]
        vc = row.cells[1]
        set_cell_bg(kc, '1E293B')
        set_cell_bg(vc, 'F8F5EE')
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.size = Pt(10)
        kr.font.color.rgb = C_GOLD
        kr.font.name = 'Calibri'
        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(10)
        vr.font.color.rgb = C_DARK_TEXT
        vr.font.name = 'Calibri'

    page_break(doc)

# ── Section 1: Executive Summary ──
def section_executive_summary(doc):
    h1(doc, '1. Executive Summary', C_PRIMARY)
    add_hr(doc, 'D4AF37')
    body(doc,
        'WebHarmonium is a zero-investment, content-first web platform designed to become the leading free '
        'online harmonium learning destination for South Asia and the global Pakistani/Indian diaspora. '
        'The platform combines an interactive virtual harmonium instrument with structured educational '
        'content, programmatic SEO, and Google AdSense monetization.')
    body(doc,
        'By leveraging free hosting (Netlify), free version control (GitHub), and organic search traffic, '
        'WebHarmonium can achieve sustainable AdSense revenue within 3–4 months of launch — entirely '
        'without any upfront capital investment.')

    h2(doc, 'Key Value Propositions')
    bullets = [
        ('🎵 Virtual Harmonium', 'Authentic multi-reed sound synthesis via Web Audio API. No downloads, no login.'),
        ('📚 Content Engine', '25 articles at launch scaling to 3,000+ programmatic SEO pages.'),
        ('💰 Zero Cost Launch', '$0 hosting on Netlify + GitHub. Domain purchase optional.'),
        ('📈 Organic Growth', 'Target 10,000+ daily visitors within 12 months via SEO.'),
        ('💵 AdSense Revenue', 'Monetize traffic with Google AdSense. Add affiliate and premium later.'),
    ]
    for emoji_title, desc in bullets:
        t = doc.add_table(1, 2)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        c0, c1 = t.rows[0].cells[0], t.rows[0].cells[1]
        set_cell_bg(c0, '1E293B')
        set_cell_bg(c1, 'F8F5EE')
        c0.width = Inches(1.8)
        r0 = c0.paragraphs[0].add_run(emoji_title)
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = C_GOLD; r0.font.name = 'Calibri'
        r1 = c1.paragraphs[0].add_run(desc)
        r1.font.size = Pt(10); r1.font.color.rgb = C_DARK_TEXT; r1.font.name = 'Calibri'
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    page_break(doc)

# ── Section 2: Brand Style Guide ──
def section_brand(doc):
    h1(doc, '2. Brand Style Guide', C_PRIMARY)
    add_hr(doc, 'D4AF37')
    body(doc, 'The WebHarmonium brand is built around a premium, educational, and culturally rich aesthetic. '
         'The design language combines the deep, resonant warmth of Indian classical music with a modern, '
         'clean digital interface.')

    h2(doc, '2.1 Brand Positioning')
    body(doc, 'Modern · Clean · Educational · Premium · South Asian Heritage')

    h2(doc, '2.2 Typography')
    body(doc, 'Careful typographic choices create a premium feel while maintaining high readability across all devices.')

    t = doc.add_table(4, 3)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ['Role', 'Font Family', 'Weight / Size']
    row = t.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(row.cells[i], '0F172A')
        r = row.cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'

    type_rows = [
        ('Headings (H1–H3)', 'Poppins', 'Bold 700/800 · 28–52px · Letter-spacing -0.02em'),
        ('Body Text', 'Inter', 'Regular 400 · 16–18px · Line-height 1.7'),
        ('Alternative Pair', 'Montserrat + Open Sans', '700 Headings · 400 Body · Same visual weight'),
    ]
    for i, (role, font, spec) in enumerate(type_rows):
        row = t.rows[i+1]
        vals = [role, font, spec]
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(vals):
            set_cell_bg(row.cells[j], bg)
            r = row.cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 1: r.bold = True

    h2(doc, '2.3 Typographic Scale')
    scale_data = [
        ('H1 — Hero', '42–52px / clamp(36px, 5vw, 52px)', 'Poppins 800', 'Page titles, hero headlines'),
        ('H2 — Section', '32–38px / clamp(28px, 3.5vw, 38px)', 'Poppins 700', 'Major section titles'),
        ('H3 — Subsection', '24–28px / clamp(20px, 2.5vw, 28px)', 'Poppins 600', 'Card titles, subsections'),
        ('H4 — Card Title', '18px', 'Poppins 600', 'Feature headings'),
        ('Body', '16–18px / clamp(15px, 1.5vw, 17px)', 'Inter 400', 'Article text, descriptions'),
        ('Lead', '17–20px', 'Inter 400', 'Hero subtext, intro paragraphs'),
        ('Small / Meta', '11–13px', 'Inter 500', 'Labels, tags, captions'),
    ]
    t2 = doc.add_table(len(scale_data)+1, 4)
    t2.style = 'Table Grid'
    for i, col in enumerate(['Type', 'Size', 'Font', 'Usage']):
        set_cell_bg(t2.rows[0].cells[i], '1E293B')
        r = t2.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_data in enumerate(scale_data):
        bg = 'FAFAFA' if i % 2 == 0 else 'F8F5EE'
        for j, val in enumerate(row_data):
            set_cell_bg(t2.rows[i+1].cells[j], bg)
            r = t2.rows[i+1].cells[j].paragraphs[0].add_run(val)
            r.font.size = Pt(9); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 0: r.bold = True

    h2(doc, '2.4 Color Palette')
    body(doc, 'The palette is inspired by the rich, deep tones of a traditional wooden harmonium paired with '
         'the golden shimmer of its brass keys — modernized for a premium digital experience.')

    colors = [
        ('#0F172A', 'Primary / Background', 'Deep Navy', 'Page backgrounds, navbar'),
        ('#1E293B', 'Secondary / Surface', 'Dark Slate', 'Cards, panels, footer'),
        ('#D4AF37', 'Accent Gold', 'Metallic Gold', 'Key highlights, headings, CTAs'),
        ('#F8F5EE', 'Accent Cream', 'Warm Cream', 'Text on dark, soft backgrounds'),
        ('#22C55E', 'CTA / Success', 'Emerald Green', 'Primary buttons, active states'),
        ('#94A3B8', 'Muted Text', 'Slate Grey', 'Body text, descriptions'),
        ('#F1F5F9', 'Light Text', 'Off-White', 'Primary text on dark backgrounds'),
    ]
    t3 = doc.add_table(len(colors)+1, 4)
    t3.style = 'Table Grid'
    for i, col in enumerate(['Hex Code', 'Role', 'Name', 'Usage']):
        set_cell_bg(t3.rows[0].cells[i], '0F172A')
        r = t3.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'

    for i, (hex_c, role, name, usage) in enumerate(colors):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        row = t3.rows[i+1]
        for j, val in enumerate([hex_c, role, name, usage]):
            set_cell_bg(row.cells[j], bg)
            r = row.cells[j].paragraphs[0].add_run(val)
            r.font.size = Pt(9); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 0:
                r.bold = True
                r.font.color.rgb = C_GOLD

    h2(doc, '2.5 Design Principles')
    principles = [
        ('Glassmorphism', 'Cards use blur + semi-transparent backgrounds for a modern, premium depth effect.'),
        ('Gold Gradients', 'Text and borders use gradient gold to evoke the warmth of physical harmonium keys.'),
        ('Micro-Animations', 'Hover effects, fade-in on scroll, key press animations create a living interface.'),
        ('Dark Mode First', 'Deep navy primary ensures comfortable long practice sessions in any lighting.'),
        ('Mobile First', 'Responsive breakpoints at 480px, 768px, 1024px ensure great mobile experience.'),
    ]
    for p_name, p_desc in principles:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        r1 = p.add_run(f'{p_name}: ')
        r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = C_GOLD; r1.font.name = 'Calibri'
        r2 = p.add_run(p_desc)
        r2.font.size = Pt(11); r2.font.color.rgb = C_DARK_TEXT; r2.font.name = 'Calibri'

    page_break(doc)

# ── Section 3: Technical Architecture ──
def section_tech(doc):
    h1(doc, '3. Technical Architecture', C_PRIMARY)
    add_hr(doc, 'D4AF37')
    body(doc, 'WebHarmonium is built on a JAMstack (JavaScript, APIs, Markup) architecture for maximum '
         'performance, security, and zero hosting cost. The entire platform runs on static files '
         'served by Netlify\'s global CDN.')

    h2(doc, '3.1 Technology Stack')
    stack = [
        ('Frontend', 'HTML5 · CSS3 · Vanilla JavaScript', 'Zero framework overhead, fast load times'),
        ('Audio Engine', 'Web Audio API', 'Real-time reed synthesis, no external audio files needed'),
        ('Styling', 'Vanilla CSS with CSS Custom Properties', 'Design tokens, no preprocessor overhead'),
        ('Fonts', 'Google Fonts (Poppins + Inter)', 'CDN-served, optimized typography'),
        ('Hosting', 'Netlify Free Tier', 'Global CDN, HTTPS, 100GB bandwidth/month'),
        ('Repository', 'GitHub', 'Version control, CI/CD pipeline'),
        ('Analytics', 'Google Analytics 4', 'Traffic tracking, audience insights'),
        ('Search', 'Google Search Console', 'Indexing, keyword performance, sitemap submission'),
        ('Monetization', 'Google AdSense', 'Display advertising revenue'),
        ('Future: Auth', 'Firebase Free Tier', 'User accounts, progress saving (Phase 3)'),
    ]
    t = doc.add_table(len(stack)+1, 3)
    t.style = 'Table Grid'
    for i, col in enumerate(['Layer', 'Technology', 'Rationale']):
        set_cell_bg(t.rows[0].cells[i], '0F172A')
        r = t.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, (layer, tech, reason) in enumerate(stack):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate([layer, tech, reason]):
            set_cell_bg(t.rows[i+1].cells[j], bg)
            r = t.rows[i+1].cells[j].paragraphs[0].add_run(val)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 0: r.bold = True

    h2(doc, '3.2 Architecture Flow')
    body(doc, 'User → Netlify CDN → Static HTML/CSS/JS → Web Audio API → Sound Output')
    body(doc, 'User → Google Search → WebHarmonium Page → AdSense Ad Display → Revenue')

    h2(doc, '3.3 File Structure')
    files = [
        ('index.html', 'Main landing page with virtual harmonium'),
        ('index.css', 'Complete design system and component styles'),
        ('harmonium.js', 'Web Audio API synth engine + keyboard handlers'),
        ('articles.json', 'SEO article metadata for programmatic content'),
        ('about.html', 'About page — mission, heritage, roadmap preview'),
        ('contact.html', 'Contact form — Netlify Forms ready'),
        ('privacy.html', 'Privacy Policy — AdSense compliance requirement'),
        ('terms.html', 'Terms of Service — AdSense compliance requirement'),
        ('sitemap.xml', '30-page sitemap for Google Search Console'),
        ('robots.txt', 'Crawler directives and sitemap pointer'),
        ('netlify.toml', 'Security headers, cache control, redirects'),
    ]
    for fname, desc in files:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        r1 = p.add_run(f'{fname}')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = C_GOLD
        r1.font.name = 'Courier New'
        r2 = p.add_run(f'  —  {desc}')
        r2.font.size = Pt(10); r2.font.color.rgb = C_DARK_TEXT; r2.font.name = 'Calibri'

    h2(doc, '3.4 Domain Strategy')
    body(doc, 'Phase 1 Launch (Free):')
    bullet(doc, 'webharmonium.netlify.app — Deploy immediately, zero cost')
    body(doc, 'After First Revenue (Recommended Domains):')
    bullet(doc, 'webharmonium.pk — Pakistani market focus (~PKR 2,000/year)')
    bullet(doc, 'learnharmonium.com — Broader international appeal')
    bullet(doc, 'harmoniumnotes.com — High search-intent, notes-focused')

    page_break(doc)

# ── Section 4: Phase 1 Plan ──
def section_phase1(doc):
    h1(doc, '4. Phase 1: MVP Launch (Weeks 1–2)', C_PRIMARY)
    add_hr(doc, '22C55E')
    body(doc, 'Goal: Launch a production-ready, fully SEO-optimized harmonium platform and get indexed by Google '
         'within 14 days. Submit AdSense application by end of Week 2.')

    h2(doc, '4.1 Core Features Built')
    feats = [
        '✅ Virtual Harmonium with 36-key, 3-octave keyboard (C3–B5)',
        '✅ Authentic multi-reed Web Audio API synthesis (Bass + Male + Treble stops)',
        '✅ Computer keyboard, mouse, and multi-touch support',
        '✅ Record & download sessions (WebM audio)',
        '✅ Sa Re Ga Ma sargam labels on all keys',
        '✅ Bellows animation with pressure simulation',
        '✅ Note visualizer showing notes as they are played',
        '✅ Volume control slider',
        '✅ Octave shift control (±3 octaves)',
    ]
    for f in feats:
        bullet(doc, f)

    h2(doc, '4.2 Pages Launched')
    pages = [
        ('index.html', 'Home + Virtual Harmonium', 'Priority 1.0'),
        ('about.html', 'About WebHarmonium', 'Priority 0.6'),
        ('contact.html', 'Contact Form', 'Priority 0.5'),
        ('privacy.html', 'Privacy Policy (AdSense Required)', 'Priority 0.3'),
        ('terms.html', 'Terms of Service (AdSense Required)', 'Priority 0.3'),
    ]
    t = doc.add_table(len(pages)+1, 3)
    t.style = 'Table Grid'
    for i, col in enumerate(['File', 'Description', 'Sitemap Priority']):
        set_cell_bg(t.rows[0].cells[i], '1E293B')
        r = t.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(pages):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row_d):
            set_cell_bg(t.rows[i+1].cells[j], bg)
            r = t.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 0: r.font.name = 'Courier New'; r.bold = True

    h2(doc, '4.3 SEO Articles (25 Seeded)')
    categories = {
        'Beginner Guides (7)': [
            'Learn Harmonium Online: Complete Beginner\'s Guide',
            'Harmonium Notes for Beginners: Sa Re Ga Ma Pa',
            'Harmonium Basics: Understanding the Instrument',
            'Harmonium Learning Guide: 30-Day Plan',
            'How to Use Bellows on Harmonium',
            'Harmonium Keyboard Layout Explained',
            'Virtual Harmonium vs Real Harmonium',
        ],
        'Practice & Technique (5)': [
            'Sa Re Ga Ma Practice: Daily Exercises',
            'Best Harmonium Exercises for Beginners',
            'Alankar Practice on Harmonium',
            'Harmonium Finger Exercises for Speed',
            'Understanding Harmonium Stops (Reeds)',
        ],
        'Raag Library (4)': [
            'How to Play Raag Yaman on Harmonium',
            'Raag Bhairav: Notes, Rules & Practice',
            'Raag Kafi: A Complete Guide',
            'Raag Bhopali: The Pentatonic Raag',
        ],
        'Songs & Naat (5)': [
            'Tajdar-e-Haram Harmonium Notes',
            'Lab Pe Aati Hai Dua Harmonium Notes',
            'Naat Harmonium Notes for Beginners',
            'Hamd Harmonium Notes Collection',
            'Playing Qawwali on Harmonium: A Guide',
        ],
        'Theory (4)': [
            'C Major Scale on Harmonium',
            'D Major Scale on Harmonium',
            'Indian Classical Music Theory for Harmonium',
            'How to Tune a Harmonium',
        ],
    }
    for cat, articles in categories.items():
        h3(doc, cat, C_SECONDARY)
        for a in articles:
            bullet(doc, a)

    h2(doc, '4.4 Analytics & Technical SEO Checklist')
    checklist = [
        '☐ Create Google Analytics 4 property → Add GA4 ID to index.html',
        '☐ Verify site in Google Search Console',
        '☐ Submit sitemap.xml to Search Console',
        '☐ Push to GitHub repository',
        '☐ Connect GitHub to Netlify for auto-deploy',
        '☐ Confirm HTTPS is active on Netlify',
        '☐ Test robots.txt at /robots.txt',
        '☐ Test sitemap at /sitemap.xml',
        '☐ Verify Open Graph tags with Facebook Debugger',
        '☐ Run Lighthouse audit (target: 90+ Performance, 100 SEO)',
        '☐ Submit AdSense application',
    ]
    for item in checklist:
        bullet(doc, item)

    h2(doc, '4.5 Phase 1 KPIs')
    kpis = [
        ('Indexed Pages', '30', '30 pages in sitemap.xml'),
        ('Daily Visitors', '100', 'Organic + social sharing'),
        ('Bounce Rate', '< 60%', 'Virtual harmonium keeps users engaged'),
        ('AdSense', 'Applied', 'Privacy + Terms pages enable application'),
    ]
    t2 = doc.add_table(len(kpis)+1, 3)
    t2.style = 'Table Grid'
    for i, col in enumerate(['KPI', 'Target', 'Notes']):
        set_cell_bg(t2.rows[0].cells[i], '0F172A')
        r = t2.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(kpis):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row_d):
            set_cell_bg(t2.rows[i+1].cells[j], bg)
            r = t2.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 1: r.bold = True; r.font.color.rgb = C_CTA

    page_break(doc)

# ── Section 5: Phase 2 ──
def section_phase2(doc):
    h1(doc, '5. Phase 2: SEO Growth (Months 2–4)', C_PRIMARY)
    add_hr(doc, 'D4AF37')
    body(doc, 'Goal: Scale organic traffic through programmatic SEO. Achieve Google AdSense approval and '
         'generate first revenue. Target 300+ indexed pages and 1,000 daily visitors.')

    h2(doc, '5.1 Content Clusters & Programmatic SEO')
    clusters = [
        ('Notes Cluster', '/notes/[note]', ['sa', 're', 'ga', 'ma', 'pa', 'dha', 'ni'], 7),
        ('Scales Cluster', '/scales/[scale]', ['c-major', 'd-major', 'e-major', 'g-major', 'a-minor', 'b-flat', 'f-major'], 7),
        ('Raags Cluster', '/raag/[raag]', ['yaman', 'bhairav', 'kafi', 'bhopali', 'marwa', 'todi', 'durga', 'desh', 'bhairavi', 'khamaj'], 50),
        ('Songs Cluster', '/song/[slug]', ['tajdar-e-haram', 'lab-pe-aati-hai-dua', 'hamd-notes', 'naat-collection', 'qawwali-pack'], 200),
    ]
    t = doc.add_table(len(clusters)+1, 4)
    t.style = 'Table Grid'
    for i, col in enumerate(['Cluster', 'URL Pattern', 'Examples', 'Pages']):
        set_cell_bg(t.rows[0].cells[i], '0F172A')
        r = t.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, (name, pattern, examples, pages) in enumerate(clusters):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        row_vals = [name, pattern, ', '.join(examples[:3]) + '...', f'{pages}+']
        for j, v in enumerate(row_vals):
            set_cell_bg(t.rows[i+1].cells[j], bg)
            r = t.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 1: r.font.name = 'Courier New'; r.font.color.rgb = C_GOLD
            if j == 3: r.bold = True; r.font.color.rgb = C_CTA

    h2(doc, '5.2 New Features: Phase 2')
    bullet(doc, 'Recording Feature — Play, record, download as WebM audio')
    bullet(doc, 'Practice Mode — Beginner, Intermediate, Advanced difficulty levels')
    bullet(doc, 'Highlighted Scale Display — Show scale notes on keyboard')
    bullet(doc, 'JSON-driven Programmatic Pages — One template, thousands of pages')
    bullet(doc, 'Mobile PWA Support — Add to Home Screen capability')
    bullet(doc, 'Social Sharing — Share recordings and article links')

    h2(doc, '5.3 Phase 2 KPIs')
    kpis2 = [
        ('Indexed Pages', '300+', 'Programmatic notes/raag/song pages'),
        ('Daily Visitors', '1,000', 'Primarily organic Google traffic'),
        ('AdSense Status', 'Approved', 'Revenue generation begins'),
        ('Monthly Revenue', '$10–50', 'AdSense estimate at 1K/day'),
    ]
    t2 = doc.add_table(len(kpis2)+1, 3)
    t2.style = 'Table Grid'
    for i, col in enumerate(['KPI', 'Target', 'Notes']):
        set_cell_bg(t2.rows[0].cells[i], '1E293B')
        r = t2.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(kpis2):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row_d):
            set_cell_bg(t2.rows[i+1].cells[j], bg)
            r = t2.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 1: r.bold = True; r.font.color.rgb = C_CTA

    page_break(doc)

# ── Section 6: Phase 3 ──
def section_phase3(doc):
    h1(doc, '6. Phase 3: Authority Platform (Months 5–12)', C_PRIMARY)
    add_hr(doc, 'D4AF37')
    body(doc, 'Goal: Become the #1 harmonium learning platform globally. 2,000–5,000 indexed pages, '
         '10,000+ daily visitors, multiple revenue streams generating stable monthly income.')

    h2(doc, '6.1 Advanced Features')
    adv = [
        ('Interactive Lessons', 'Structured video + exercise courses: Basic → Scale → Song → Raag → Advanced'),
        ('User Accounts', 'Firebase Free Tier: Save progress, favorites, practice history'),
        ('AI Features', 'AI Song Notes Generator, AI Scale Recommender, AI Practice Assistant (free APIs)'),
        ('Community Section', 'Users upload notes, share tutorials, comment on lessons'),
        ('Lesson Gamification', 'Streaks, badges, progress bars — increase retention'),
    ]
    for name, desc in adv:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{name}: ')
        r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = C_GOLD; r1.font.name = 'Calibri'
        r2 = p.add_run(desc)
        r2.font.size = Pt(11); r2.font.color.rgb = C_DARK_TEXT; r2.font.name = 'Calibri'

    h2(doc, '6.2 SEO Scale')
    seo_scale = [
        ('Songs', '500+ song pages', 'Popular to long-tail search queries'),
        ('Raags', '100+ raag pages', 'Classical music reference library'),
        ('Notes', '100+ note pages', '/notes/sa, /notes/re, etc.'),
        ('Tutorials', '500+ tutorial pages', 'Step-by-step video-linked guides'),
        ('FAQs', '1,000+ FAQ pages', 'Voice search & featured snippets'),
    ]
    t = doc.add_table(len(seo_scale)+1, 3)
    t.style = 'Table Grid'
    for i, col in enumerate(['Content Type', 'Volume', 'SEO Strategy']):
        set_cell_bg(t.rows[0].cells[i], '0F172A')
        r = t.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(seo_scale):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row_d):
            set_cell_bg(t.rows[i+1].cells[j], bg)
            r = t.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 1: r.bold = True; r.font.color.rgb = C_CTA

    h2(doc, '6.3 Monetization Strategy')
    revenue = [
        ('Google AdSense', 'Primary', 'Display ads on all content pages. At 10K/day with $1 RPM = $300/month'),
        ('Affiliate Marketing', 'Secondary', 'Harmonium products, music books, online courses on Amazon/Flipkart'),
        ('Premium Content', 'Tertiary', 'Advanced lessons, song packs, practice exercises ($5–15/month)'),
        ('Donations', 'Optional', 'Ko-fi / Patreon for community support'),
        ('Sponsored Content', 'Future', 'Music schools, instrument retailers sponsor articles'),
    ]
    t2 = doc.add_table(len(revenue)+1, 3)
    t2.style = 'Table Grid'
    for i, col in enumerate(['Revenue Source', 'Priority', 'Estimated Potential']):
        set_cell_bg(t2.rows[0].cells[i], '0F172A')
        r = t2.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, (src, pri, est) in enumerate(revenue):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate([src, pri, est]):
            set_cell_bg(t2.rows[i+1].cells[j], bg)
            r = t2.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 0: r.bold = True

    page_break(doc)

# ── Section 7: Success Metrics ──
def section_metrics(doc):
    h1(doc, '7. Success Metrics & Growth Milestones', C_PRIMARY)
    add_hr(doc, 'D4AF37')

    milestones = [
        ('Month 1', '30', '100', 'AdSense Applied', '$0'),
        ('Month 2', '100', '300', 'AdSense Review', '$0–5'),
        ('Month 3', '300', '1,000', 'AdSense Approved', '$10–50'),
        ('Month 6', '1,000', '5,000', 'Stable Revenue', '$100–500'),
        ('Month 12', '3,000+', '10,000–50,000', 'Authority Platform', '$500–3,000+'),
    ]
    t = doc.add_table(len(milestones)+1, 5)
    t.style = 'Table Grid'
    for i, col in enumerate(['Month', 'Pages', 'Daily Visitors', 'Milestone', 'Est. Monthly Revenue']):
        set_cell_bg(t.rows[0].cells[i], '0F172A')
        r = t.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(milestones):
        bg = '1E293B' if i % 2 == 0 else '0F172A'
        text_colors = [C_GOLD, C_CTA, C_CREAM, C_GREY, C_CTA]
        for j, v in enumerate(row_d):
            set_cell_bg(t.rows[i+1].cells[j], 'F8F5EE' if i % 2 == 0 else 'FFFFFF')
            r = t.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j in [1, 2, 4]: r.bold = True; r.font.color.rgb = C_CTA
            if j == 0: r.bold = True

    h2(doc, 'AdSense Revenue Projections')
    body(doc, 'Estimated based on music education niche RPM of $0.80–$2.00:')
    proj = [
        ('1,000 visitors/day', '30,000 page views/month', '$24–60/month'),
        ('5,000 visitors/day', '150,000 page views/month', '$120–300/month'),
        ('10,000 visitors/day', '300,000 page views/month', '$240–600/month'),
        ('50,000 visitors/day', '1,500,000 page views/month', '$1,200–3,000/month'),
    ]
    t2 = doc.add_table(len(proj)+1, 3)
    t2.style = 'Table Grid'
    for i, col in enumerate(['Traffic', 'Monthly Page Views', 'Est. AdSense Revenue']):
        set_cell_bg(t2.rows[0].cells[i], '1E293B')
        r = t2.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(proj):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row_d):
            set_cell_bg(t2.rows[i+1].cells[j], bg)
            r = t2.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 2: r.bold = True; r.font.color.rgb = C_CTA

    page_break(doc)

# ── Section 8: Launch Checklist ──
def section_launch(doc):
    h1(doc, '8. Launch Checklist', C_PRIMARY)
    add_hr(doc, '22C55E')

    tasks = {
        'GitHub Setup': [
            'Create GitHub account (free)',
            'Create repository: webharmonium',
            'Push all project files to main branch',
            'Add .gitignore for node_modules (future use)',
        ],
        'Netlify Deployment': [
            'Create Netlify account (free)',
            'Connect GitHub repository to Netlify',
            'Configure build settings (static site, no build command)',
            'Verify site is live at webharmonium.netlify.app',
            'Confirm HTTPS is active',
            'Test netlify.toml headers are working',
        ],
        'Google Analytics': [
            'Create Google Analytics 4 property',
            'Copy Measurement ID (G-XXXXXXXXXX)',
            'Add GA4 script to all HTML pages',
            'Verify data is flowing in GA4 dashboard',
        ],
        'Google Search Console': [
            'Add webharmonium.netlify.app as property',
            'Verify ownership via HTML tag or DNS',
            'Submit sitemap.xml',
            'Request indexing for homepage',
            'Monitor Coverage report for indexing errors',
        ],
        'AdSense Application': [
            'Confirm Privacy Policy page is live and accessible',
            'Confirm Terms of Service page is live',
            'Ensure at least 20-30 quality pages are indexed',
            'Apply at adsense.google.com',
            'Add AdSense code to site after approval',
        ],
    }

    for section_title, items in tasks.items():
        h2(doc, section_title)
        for item in items:
            bullet(doc, f'☐  {item}')

    page_break(doc)

# ── Section 9: Appendix ──
def section_appendix(doc):
    h1(doc, '9. Appendix: SEO Keywords & Content Calendar', C_PRIMARY)
    add_hr(doc, 'D4AF37')

    h2(doc, '9.1 High-Priority Keywords (Phase 1)')
    kws = [
        ('learn harmonium online', 'High', 'Navigational / Informational', 'Beginner Guide Article'),
        ('virtual harmonium', 'High', 'Transactional', 'Homepage'),
        ('harmonium notes for beginners', 'High', 'Informational', 'Notes Article'),
        ('sa re ga ma harmonium', 'High', 'Informational', 'Practice Article'),
        ('harmonium keyboard layout', 'Medium', 'Informational', 'Layout Article'),
        ('raag yaman notes', 'Medium', 'Informational', 'Raag Yaman Article'),
        ('tajdar e haram harmonium notes', 'Medium', 'Informational', 'Song Notes Article'),
        ('online harmonium player', 'Medium', 'Transactional', 'Homepage'),
        ('harmonium online free', 'High', 'Transactional', 'Homepage'),
        ('harmonium for beginners', 'High', 'Informational', 'Beginner Guides'),
    ]
    t = doc.add_table(len(kws)+1, 4)
    t.style = 'Table Grid'
    for i, col in enumerate(['Keyword', 'Volume', 'Intent', 'Target Page']):
        set_cell_bg(t.rows[0].cells[i], '0F172A')
        r = t.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(kws):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row_d):
            set_cell_bg(t.rows[i+1].cells[j], bg)
            r = t.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(9); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 1:
                r.bold = True
                r.font.color.rgb = C_CTA if v == 'High' else C_GOLD

    h2(doc, '9.2 Month 1 Content Calendar')
    cal = [
        ('Week 1', 'Launch homepage + virtual harmonium + 5 core articles'),
        ('Week 2', 'Publish 20 more articles + submit sitemap to Search Console'),
        ('Week 3', 'Social sharing (Reddit, Facebook music groups) + AdSense application'),
        ('Week 4', 'Monitor Search Console for indexing + respond to user feedback'),
    ]
    for week, task in cal:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{week}: ')
        r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = C_GOLD; r1.font.name = 'Calibri'
        r2 = p.add_run(task)
        r2.font.size = Pt(11); r2.font.color.rgb = C_DARK_TEXT; r2.font.name = 'Calibri'

    h2(doc, '9.3 Resources & Tools (All Free)')
    resources = [
        ('Netlify', 'netlify.com', 'Hosting, SSL, CDN'),
        ('GitHub', 'github.com', 'Version control, CI/CD'),
        ('Google Analytics', 'analytics.google.com', 'Traffic analytics'),
        ('Google Search Console', 'search.google.com/search-console', 'SEO monitoring'),
        ('Google AdSense', 'adsense.google.com', 'Advertising revenue'),
        ('PageSpeed Insights', 'pagespeed.web.dev', 'Performance testing'),
        ('Google Fonts', 'fonts.google.com', 'Poppins + Inter fonts'),
        ('Favicon Generator', 'favicon.io', 'Custom favicon creation'),
    ]
    t2 = doc.add_table(len(resources)+1, 3)
    t2.style = 'Table Grid'
    for i, col in enumerate(['Tool', 'URL', 'Purpose']):
        set_cell_bg(t2.rows[0].cells[i], '1E293B')
        r = t2.rows[0].cells[i].paragraphs[0].add_run(col)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'
    for i, row_d in enumerate(resources):
        bg = 'F8F5EE' if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row_d):
            set_cell_bg(t2.rows[i+1].cells[j], bg)
            r = t2.rows[i+1].cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10); r.font.color.rgb = C_DARK_TEXT; r.font.name = 'Calibri'
            if j == 0: r.bold = True
            if j == 1: r.font.color.rgb = C_GOLD

    # Final page: sign-off
    page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run('🎵')
    r.font.size = Pt(48)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('WebHarmonium')
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = C_PRIMARY; r.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Build the music. Build the audience. Build the revenue.')
    r.font.size = Pt(14); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'; r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('webharmonium.netlify.app')
    r.font.size = Pt(12); r.font.color.rgb = C_GREY; r.font.name = 'Calibri'

    add_hr(doc, 'D4AF37')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('© 2024 WebHarmonium — Made with ♥ for music lovers of South Asia')
    r.font.size = Pt(10); r.font.color.rgb = C_GREY; r.font.name = 'Calibri'

# ── Main ──
def generate():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    make_cover(doc)
    section_executive_summary(doc)
    section_brand(doc)
    section_tech(doc)
    section_phase1(doc)
    section_phase2(doc)
    section_phase3(doc)
    section_metrics(doc)
    section_launch(doc)
    section_appendix(doc)

    output = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'WebHarmonium_Roadmap.docx')
    output = os.path.normpath(output)
    doc.save(output)
    print(f"\n✅ Document generated successfully!")
    print(f"📄 Saved to: {output}")
    print(f"\nDocument contains:")
    print("  • Cover page with brand identity")
    print("  • Executive Summary")
    print("  • Brand Style Guide (typography, colors, design principles)")
    print("  • Technical Architecture")
    print("  • Phase 1: MVP Launch (Weeks 1–2)")
    print("  • Phase 2: SEO Growth (Months 2–4)")
    print("  • Phase 3: Authority Platform (Months 5–12)")
    print("  • Success Metrics & Revenue Projections")
    print("  • Launch Checklist")
    print("  • Appendix: Keywords & Content Calendar")

if __name__ == '__main__':
    generate()
